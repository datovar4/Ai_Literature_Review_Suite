"""
Created on Sat May 13 19:00:28 2023

@author: David Tovar
"""
import os
import openai
from tkinter import Tk, filedialog, simpledialog, Label, Entry
import sys
import pandas as pd
from docx import Document
from sklearn.neighbors import NearestNeighbors
import tensorflow_hub as hub
import numpy as np
from sklearn.cluster import KMeans
from math import ceil
from os.path import dirname
import requests
import time
import argparse
from api_key_checker import check_api_key

# Set your OpenAI API Key here
check_api_key("OpenAI_API.txt", "OpenAI")
with open("OpenAI_API.txt", "r") as f:
    openai_api_key = f.read().strip()
    
openai.api_key = openai_api_key


class SemanticSearch:
    def __init__(self):
        self.use = hub.load('https://tfhub.dev/google/universal-sentence-encoder/4')
        self.fitted = False

    def fit(self, data, batch=1000, n_neighbors=5, entries_per_cluster=10):
        self.data = data
        self.embeddings, self.clusters = self.get_text_embedding(data, batch=batch, entries_per_cluster=entries_per_cluster)
        n_neighbors = min(n_neighbors, len(self.embeddings))
        self.nn = NearestNeighbors(n_neighbors=n_neighbors)
        self.nn.fit(self.embeddings)
        self.fitted = True

    def __call__(self, text, return_data=True):
        inp_emb = self.use([text])
        neighbors = self.nn.kneighbors(inp_emb, return_distance=False)[0]

        if return_data:
            return [self.data[i] for i in neighbors]
        else:
            return neighbors

    def get_text_embedding(self, texts, batch=1000, entries_per_cluster=10):
        embeddings = []
        for i in range(0, len(texts), batch):
            text_batch = texts[i:(i+batch)]
            emb_batch = self.use(text_batch)
            embeddings.append(emb_batch)
        embeddings = np.vstack(embeddings)

        # Determine the number of clusters dynamically based on the total number of entries
        n_clusters = ceil(len(texts) / entries_per_cluster)

        # Perform clustering on the embeddings
        kmeans = KMeans(n_clusters=n_clusters, random_state=3)
        clusters = kmeans.fit_predict(embeddings)

        # Build a list of lists where each sub-list contains the texts for a particular cluster
        text_clusters = [[] for _ in range(n_clusters)]
        for text, cluster in zip(texts, clusters):
            text_clusters[cluster].append(text)

        # Redistribute texts across clusters for a more balanced distribution
        uneven_clusters = [cluster for cluster in text_clusters if len(cluster) > entries_per_cluster]
        for cluster in uneven_clusters:
            num_excess_entries = len(cluster) - entries_per_cluster
            for _ in range(num_excess_entries):
                smallest_cluster = min(text_clusters, key=len)
                text = cluster.pop(0)
                smallest_cluster.append(text)

        return embeddings, text_clusters


def generate_text(prompt, engine="gpt-4o", max_tokens=300):
    max_attempts = 5
    sleep_time = 30  # Number of seconds to wait between attempts

    for attempt in range(max_attempts):
        try:
            if engine == "gpt-4o":
                completions = openai.ChatCompletion.create(
                    model=engine,
                    messages=[
                        {"role": "system", "content": "You are an expert scientific writer who writes review articles for Nature Neuroscience Reviews."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=max_tokens,
                    temperature=0.7
                )
            else:
                completions = openai.Completion.create(
                    engine=engine,
                    prompt=prompt,
                    max_tokens=max_tokens,
                    n=1,
                    stop=None,
                    temperature=0.7,
                )
            message = completions.choices[0].text.strip() if engine != "gpt-4o" else completions.choices[0].message['content'].strip()
            return message

        except requests.exceptions.RequestException as e:
            if 'rate limit' in str(e):
                if attempt + 1 == max_attempts:
                    raise e  # If this was the last attempt, raise the exception
                else:
                    time.sleep(sleep_time)  # Wait for a while before trying again
            else:
                raise e  # If this was not a rate limit error, raise the exception

def generate_answer(cluster_texts):
    question = "What are the main themes and what are some similarities and differences of the studies? Place some emphasis on main study findings and methods used as well as model architectures if applicable"

    prompt = ""
    for c in cluster_texts:
        prompt += c + '\n\n'

    prompt += f"In the above text, {question}. Write your response in a review style for a journal, citing the entry titles with in-text citations from the text above as your source material after each sentence.Be concise limiting your responses to a 150-200 word synthesized paragraph. Make a topic sentence that summarizes the main theme. Do not use the phrase 'in this review'. Do not make a reference to the number of articles you are summarizing. Avoid being repetitive."
    print('Calling GPT4o API for cluster summary')
    answer = generate_text(prompt, max_tokens=500)  # Increase if needed
    return answer

# Custom dialog class to input topics
class TopicsDialog(simpledialog.Dialog):
    def body(self, master):
        self.title("Input Topics for Literature Review")
        Label(master, text="Topic 1:").grid(row=0, column=0, sticky='e')
        Label(master, text="Topic 2:").grid(row=1, column=0, sticky='e')
        Label(master, text="Topic 3:").grid(row=2, column=0, sticky='e')
        Label(master, text="Main Topic:").grid(row=3, column=0, sticky='e')

        self.topic1_entry = Entry(master)
        self.topic2_entry = Entry(master)
        self.topic3_entry = Entry(master)
        self.main_topic_entry = Entry(master)

        self.topic1_entry.grid(row=0, column=1)
        self.topic2_entry.grid(row=1, column=1)
        self.topic3_entry.grid(row=2, column=1)
        self.main_topic_entry.grid(row=3, column=1)
        return self.topic1_entry  # initial focus

    def apply(self):
        self.result = {
            'topic1': self.topic1_entry.get(),
            'topic2': self.topic2_entry.get(),
            'topic3': self.topic3_entry.get(),
            'main_topic': self.main_topic_entry.get()
        }

def main(excel_file_path=None):
    # Get the Excel file path using GUI
    root = Tk()
    root.withdraw()  # Hide the main window
    root.title("Summaries Excel File")

    if excel_file_path is None:
        excel_file_path = filedialog.askopenfilename(
            title='Select Excel File with Summaries', filetypes=[('Excel files', '*.xlsx')])

    if not excel_file_path:
        print("No Excel file selected. Exiting.")
        sys.exit(1)

    # Extract base name from excel_file_path to generate Word document names
    base_name = os.path.splitext(os.path.basename(excel_file_path))[0]
    base_name = base_name.replace('question_and_answers', '')  # Remove 'question_and_answers' from the name
    doc_file_name = base_name + "_summary.docx"
    verbose_doc_file_name = base_name + "_summary_verbose.docx"


    # Read the Excel file into a DataFrame
    df_dict = pd.read_excel(excel_file_path, sheet_name=None)

    global recommender
    recommender = SemanticSearch()

    # Initialize the Word document
    if os.path.exists(doc_file_name):
        document = Document(doc_file_name)
    else:
        document = Document()

    # Initialize the verbose Word document
    verbose_doc_file_name = doc_file_name.replace('.docx', '_verbose.docx')
    if os.path.exists(verbose_doc_file_name):
        verbose_document = Document(verbose_doc_file_name)
    else:
        verbose_document = Document()

   # Initialize parameters
    entries_per_cluster = 5
    max_attempts = 5

    # Loop over all dataframes (i.e., all sheets in the Excel file)
    for sheet_name, df in df_dict.items():
        # Fit the recommender to the data
        data = df.apply(lambda row: ' '.join(str(e) for e in row), axis=1).tolist()
    
        for attempt in range(max_attempts):
            try:
                # Fit the recommender with the current entries_per_cluster
                recommender.fit(data, entries_per_cluster=entries_per_cluster)
    
                # Generate answers for each cluster of texts
                answers = []
                for cluster_texts in recommender.clusters:
                    answer = generate_answer(cluster_texts)
                    answers.append(answer)
                    document.add_paragraph(answer)
    
                print(f"Summaries generated from sheet: {sheet_name}")
                break  # If no exception occurred, break the retry loop
            except Exception as e:
                print(f"Attempt {attempt + 1} failed: {str(e)}")
                if attempt + 1 < max_attempts:
                    # Decrease the entries_per_cluster for the next attempt
                    entries_per_cluster = entries_per_cluster - 1
                    print(f"Retrying with smaller cluster size: {entries_per_cluster}")
                else:
                    print("All attempts failed. Please check the error messages.")
                    raise e  # If all attempts failed, raise the last exception

        # Save clusters in the verbose Word document
        for i, cluster_texts in enumerate(recommender.clusters):
            verbose_document.add_paragraph(f"Cluster {i} contains {len(cluster_texts)} texts:")
            for text in cluster_texts:
                verbose_document.add_paragraph(text)
            verbose_document.add_paragraph()
    
        print(f"Clusters generated from sheet: {sheet_name}")

        # After generating the summaries, prompt the user for topics
        # Use the custom dialog
        root.deiconify()  # Show the root window
        topics_dialog = TopicsDialog(root)
        root.withdraw()  # Hide the root window again

        if topics_dialog.result is None:
            print("No topics entered. Exiting.")
            sys.exit(1)
        else:
            topic1 = topics_dialog.result['topic1']
            topic2 = topics_dialog.result['topic2']
            topic3 = topics_dialog.result['topic3']
            main_topic = topics_dialog.result['main_topic']

        # Prepare the combined summary
        combined_summary = '\n\n'.join(answers)

        # Prepare the final prompt
        final_prompt = f"{combined_summary}\n\nCan you write a cohesive, detailed literature review in the style of Nature Neuroscience with APA style citations at the end of key sentences and synthesizing across studies that specifically address {topic1}, {topic2}, and {topic3}. Make sure to emphasize common methods {main_topic} and go into as much detail as possible. Do not mention the studies in sentences, only in the citations, instead focus on the ideas, synthesizing across studies. Best sentences are those that appropriately include multiple citations and interspersed throughout the paragraph, with added bonus if multiple citations are within those interspersed sentences. Include title subsections where appropriate. The review should be extensive, covering all relevant aspects, detailed descriptions, and should be cohesive with smooth transitions between sections and paragraphs. Include subsections where appropriate. Again, make sure your paragraphs are synthesizing across studies with multiple citations for statements you make."

        # Generate the final cohesive review
        print('Generating final cohesive review with GPT4o API')
        final_review = generate_text(final_prompt, max_tokens=4050)  # Increase max_tokens for longer response

        # Add the final review to the document
        document.add_heading(f"Final Cohesive Review for Sheet '{sheet_name}'", level=1)
        document.add_paragraph(final_review)

    # Save the Word document
    document_path = os.path.join(dirname(excel_file_path), doc_file_name)
    document.save(document_path)
    print(f"Summaries appended to the combined Word document: {document_path}")

    # Save the verbose Word document
    verbose_document_path = os.path.join(dirname(excel_file_path), verbose_doc_file_name)
    verbose_document.save(verbose_document_path)
    print(f"Clusters saved in the verbose Word document: {verbose_document_path}")

    # Open the combined Word document
    try:
        os.startfile(document_path)
    except AttributeError:
        # os.startfile is only available on Windows
        print("Automatic opening of documents is only supported on Windows.")

    # Open the verbose Word document
    try:
        os.startfile(verbose_document_path)
    except AttributeError:
        print("Automatic opening of documents is only supported on Windows.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--excel_file_path", type=str, help="Path to the excel file", default=None)
    args = parser.parse_args()
    main(args.excel_file_path)




   
   

   


   
